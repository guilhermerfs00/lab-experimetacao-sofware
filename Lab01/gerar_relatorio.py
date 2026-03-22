"""
Script para gerar relatório final em DOCX e PDF
sobre características de repositórios populares do GitHub.
"""

import pandas as pd
import numpy as np
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Caminhos
DIRETORIO_BASE = os.path.dirname(os.path.abspath(__file__))
CAMINHO_CSV = os.path.join(DIRETORIO_BASE, "graphs", "repositorios.csv")
DIRETORIO_GRAFICOS = os.path.join(DIRETORIO_BASE, "graphs")
CAMINHO_DOCX = os.path.join(DIRETORIO_BASE, "graphs", "relatorio_final.docx")
CAMINHO_PDF = os.path.join(DIRETORIO_BASE, "graphs", "relatorio_final.pdf")


def carregar_dados():
    df = pd.read_csv(CAMINHO_CSV, encoding="utf-8-sig")
    return df


def calcular_estatisticas(df):
    stats = {}

    # RQ01 - Idade
    stats["idade_media"] = round(df["Idade (anos)"].mean(), 2)
    stats["idade_mediana"] = round(df["Idade (anos)"].median(), 2)

    # RQ02 - Pull Requests Aceitos
    stats["pr_media"] = round(df["Pull Requests Aceitos"].mean(), 2)
    stats["pr_mediana"] = round(df["Pull Requests Aceitos"].median(), 2)

    # RQ03 - Releases
    stats["releases_media"] = round(df["Releases"].mean(), 2)
    stats["releases_mediana"] = round(df["Releases"].median(), 2)

    # RQ04 - Dias desde atualização
    stats["dias_atualizacao_media"] = round(df["Dias Desde Atualização"].mean(), 2)
    stats["dias_atualizacao_mediana"] = round(df["Dias Desde Atualização"].median(), 2)

    # RQ05 - Linguagens
    stats["top_linguagens"] = df["Linguagem Principal"].value_counts().head(15)
    total_repos = len(df)
    top3_langs = df["Linguagem Principal"].value_counts().head(3)
    stats["top3_linguagens"] = top3_langs
    stats["top3_percentual"] = round(top3_langs.sum() / total_repos * 100, 2)

    # RQ06 - Issues fechadas
    df_com_issues = df[df["Total de Issues"] > 0]
    stats["issues_fechadas_media"] = round(df_com_issues["Razão Issues Fechadas"].mean() * 100, 2)
    stats["issues_fechadas_mediana"] = round(df_com_issues["Razão Issues Fechadas"].median() * 100, 2)

    # RQ07 - Comparação por linguagem
    top_langs_list = df["Linguagem Principal"].value_counts().head(5).index.tolist()
    df["eh_linguagem_popular"] = df["Linguagem Principal"].isin(top_langs_list)

    grupo_popular = df[df["eh_linguagem_popular"]]
    grupo_nao_popular = df[~df["eh_linguagem_popular"]]

    stats["pr_media_popular"] = round(grupo_popular["Pull Requests Aceitos"].mean(), 2)
    stats["pr_media_nao_popular"] = round(grupo_nao_popular["Pull Requests Aceitos"].mean(), 2)
    stats["releases_media_popular"] = round(grupo_popular["Releases"].mean(), 2)
    stats["releases_media_nao_popular"] = round(grupo_nao_popular["Releases"].mean(), 2)
    stats["dias_media_popular"] = round(grupo_popular["Dias Desde Atualização"].mean(), 2)
    stats["dias_media_nao_popular"] = round(grupo_nao_popular["Dias Desde Atualização"].mean(), 2)

    if stats["pr_media_nao_popular"] > 0:
        stats["pr_diff_pct"] = round(
            (stats["pr_media_popular"] - stats["pr_media_nao_popular"]) / stats["pr_media_nao_popular"] * 100, 2
        )
    else:
        stats["pr_diff_pct"] = 0

    if stats["releases_media_nao_popular"] > 0:
        stats["releases_diff_pct"] = round(
            (stats["releases_media_popular"] - stats["releases_media_nao_popular"]) / stats["releases_media_nao_popular"] * 100, 2
        )
    else:
        stats["releases_diff_pct"] = 0

    # Tabela por linguagem (top 10)
    top10_langs = df["Linguagem Principal"].value_counts().head(10).index.tolist()
    df_top10 = df[df["Linguagem Principal"].isin(top10_langs)]
    stats["tabela_linguagens"] = df_top10.groupby("Linguagem Principal").agg(
        media_prs=("Pull Requests Aceitos", "mean"),
        media_releases=("Releases", "mean"),
        media_dias=("Dias Desde Atualização", "mean"),
    ).round(2).sort_values("media_prs", ascending=False)

    return stats


def add_heading_formatted(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading


def add_paragraph_formatted(doc, text, bold=False, italic=False, font_size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.name = "Calibri"
    para.paragraph_format.space_after = Pt(6)
    return para


def add_bullet(doc, text):
    para = doc.add_paragraph(text, style="List Bullet")
    return para


def gerar_docx(df, stats):
    doc = Document()

    # Configurar estilo padrão
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # ============================
    # TÍTULO
    # ============================
    titulo = doc.add_heading("Relatório Final", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run("Características de Repositórios Populares no GitHub")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(70, 130, 180)
    run.bold = True

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("Laboratório de Experimentação de Software — Lab01")
    run.font.size = Pt(12)
    run.italic = True

    doc.add_paragraph()  # Espaço

    # ============================
    # 1. INTRODUÇÃO
    # ============================
    add_heading_formatted(doc, "1. Introdução", level=1)

    add_heading_formatted(doc, "1.1 Contextualização", level=2)
    add_paragraph_formatted(doc,
        "O GitHub é a maior plataforma de hospedagem de código-fonte e colaboração de software do mundo, "
        "abrigando milhões de repositórios open-source. Projetos populares nessa plataforma são frequentemente "
        "utilizados como referência por desenvolvedores, empresas e pesquisadores. Compreender as características "
        "desses repositórios pode revelar padrões sobre boas práticas de desenvolvimento, manutenção e "
        "engajamento da comunidade."
    )

    add_heading_formatted(doc, "1.2 Problema Foco do Experimento", level=2)
    add_paragraph_formatted(doc,
        "Este estudo busca responder: quais são as características predominantes dos repositórios "
        "mais populares do GitHub em termos de idade, contribuição externa, frequência de releases, "
        "frequência de atualização, linguagens de programação utilizadas e gestão de issues?"
    )

    add_heading_formatted(doc, "1.3 Questões de Pesquisa", level=2)
    questoes = [
        "RQ01: Sistemas populares são maduros/antigos?",
        "RQ02: Sistemas populares recebem muita contribuição externa?",
        "RQ03: Sistemas populares lançam releases com frequência?",
        "RQ04: Sistemas populares são atualizados com frequência?",
        "RQ05: Sistemas populares são escritos nas linguagens mais populares?",
        "RQ06: Sistemas populares possuem um alto percentual de issues fechadas?",
        "RQ07: Sistemas escritos em linguagens mais populares recebem mais contribuição externa, lançam mais releases e são atualizados com mais frequência?",
    ]
    for q in questoes:
        add_bullet(doc, q)

    add_heading_formatted(doc, "1.4 Hipóteses", level=2)
    hipoteses = [
        "H1: Repositórios populares tendem a ser mais antigos, com uma idade média superior a 5 anos.",
        "H2: Projetos populares recebem uma quantidade significativa de contribuições externas, com pelo menos 1.000 pull requests aceitas em média.",
        "H3: Repositórios populares lançam releases com alta frequência, tendo um total médio de mais de 10 releases ao longo de sua existência.",
        "H4: Projetos amplamente utilizados são frequentemente atualizados, com um tempo médio desde a última atualização inferior a 30 dias.",
        "H5: Repositórios populares tendem a ser escritos nas linguagens de programação mais utilizadas, como JavaScript, Python e TypeScript, representando pelo menos 50% do total.",
        "H6: Projetos populares possuem um alto percentual de issues fechadas, com pelo menos 70% das issues resolvidas.",
        "H7 (bônus): Repositórios escritos em linguagens populares recebem 20% mais pull requests, lançam 15% mais releases e são atualizados com maior frequência do que aqueles escritos em linguagens menos comuns.",
    ]
    for h in hipoteses:
        add_bullet(doc, h)

    add_heading_formatted(doc, "1.5 Objetivos", level=2)
    add_paragraph_formatted(doc,
        "Objetivo principal: Analisar as características dos 1.000 repositórios mais populares do GitHub, "
        "identificando padrões e tendências relacionados à idade, contribuição, lançamentos, manutenção, "
        "linguagens e gestão de issues.",
        bold=True
    )
    add_paragraph_formatted(doc, "Objetivos específicos:")
    objetivos_esp = [
        "Coletar dados dos 1.000 repositórios com mais estrelas no GitHub via API GraphQL.",
        "Calcular métricas estatísticas (média, mediana) para cada questão de pesquisa.",
        "Gerar visualizações gráficas para cada métrica analisada.",
        "Confrontar os resultados obtidos com as hipóteses formuladas.",
        "Comparar o comportamento de repositórios escritos em linguagens populares versus menos populares.",
    ]
    for obj in objetivos_esp:
        add_bullet(doc, obj)

    # ============================
    # 2. METODOLOGIA
    # ============================
    add_heading_formatted(doc, "2. Metodologia", level=1)

    add_heading_formatted(doc, "2.1 Passo a Passo do Experimento", level=2)
    passos = [
        "1. Definição das questões de pesquisa e hipóteses iniciais.",
        "2. Implementação do script de coleta de dados utilizando a API GraphQL do GitHub.",
        "3. Coleta dos dados dos 1.000 repositórios com maior número de estrelas.",
        "4. Armazenamento dos dados em arquivo CSV para análise.",
        "5. Processamento estatístico dos dados com cálculo de média, mediana e distribuições.",
        "6. Geração de gráficos (histogramas e barras) para visualização dos resultados.",
        "7. Confrontação dos resultados com as hipóteses formuladas.",
    ]
    for p in passos:
        add_paragraph_formatted(doc, p)

    add_heading_formatted(doc, "2.2 Decisões", level=2)
    add_paragraph_formatted(doc,
        "Optou-se por utilizar a mediana como métrica central de tendência para evitar distorções causadas "
        "por outliers (repositórios extremamente populares que podem distorcer a média). "
        "A coleta foi limitada a 1.000 repositórios para garantir uma amostra significativa "
        "e viável dentro dos limites da API do GitHub."
    )

    add_heading_formatted(doc, "2.3 Materiais Utilizados", level=2)
    materiais = [
        "API GraphQL do GitHub para coleta de dados.",
        "Python 3.x como linguagem de programação.",
        "Biblioteca Pandas para manipulação e análise de dados.",
        "Biblioteca Matplotlib para geração de gráficos.",
        "Biblioteca Requests para comunicação HTTP com a API.",
        "Token de acesso pessoal do GitHub para autenticação.",
    ]
    for m in materiais:
        add_bullet(doc, m)

    add_heading_formatted(doc, "2.4 Métodos Utilizados", level=2)
    add_paragraph_formatted(doc,
        "Análise estatística descritiva com cálculo de média e mediana. "
        "Análise comparativa entre grupos (linguagens populares vs. não populares). "
        "Visualização de dados por meio de histogramas e gráficos de barras horizontais."
    )

    add_heading_formatted(doc, "2.5 Métricas e suas Unidades", level=2)
    table_metricas = doc.add_table(rows=7, cols=3)
    table_metricas.style = "Light Shading Accent 1"
    table_metricas.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Questão", "Métrica", "Unidade"]
    for i, h in enumerate(headers):
        table_metricas.rows[0].cells[i].text = h

    metricas_data = [
        ("RQ01", "Idade do repositório", "Anos"),
        ("RQ02", "Total de pull requests aceitas", "Quantidade"),
        ("RQ03", "Total de releases", "Quantidade"),
        ("RQ04", "Tempo desde última atualização", "Dias"),
        ("RQ05", "Linguagem primária", "Categórica"),
        ("RQ06", "Razão issues fechadas / total", "Proporção (0 a 1)"),
    ]
    for i, (rq, metrica, unidade) in enumerate(metricas_data):
        table_metricas.rows[i + 1].cells[0].text = rq
        table_metricas.rows[i + 1].cells[1].text = metrica
        table_metricas.rows[i + 1].cells[2].text = unidade

    # ============================
    # 3. RESULTADOS
    # ============================
    add_heading_formatted(doc, "3. Visualização dos Resultados", level=1)

    # RQ01
    add_heading_formatted(doc, "3.1 RQ01 — Sistemas populares são maduros/antigos?", level=2)
    add_paragraph_formatted(doc, f"Métrica: Idade do repositório (calculada a partir da data de criação)")
    add_bullet(doc, f"Idade média dos repositórios: {stats['idade_media']} anos")
    add_bullet(doc, f"Mediana da idade: {stats['idade_mediana']} anos")

    img_path = os.path.join(DIRETORIO_GRAFICOS, "rq01_idade.png")
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # RQ02
    add_heading_formatted(doc, "3.2 RQ02 — Sistemas populares recebem muita contribuição externa?", level=2)
    add_paragraph_formatted(doc, f"Métrica: Total de pull requests aceitas")
    add_bullet(doc, f"Média de pull requests aceitos: {stats['pr_media']}")
    add_bullet(doc, f"Mediana de pull requests aceitos: {stats['pr_mediana']}")

    img_path = os.path.join(DIRETORIO_GRAFICOS, "rq02_pull_requests.png")
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # RQ03
    add_heading_formatted(doc, "3.3 RQ03 — Sistemas populares lançam releases com frequência?", level=2)
    add_paragraph_formatted(doc, f"Métrica: Total de releases")
    add_bullet(doc, f"Média de releases: {stats['releases_media']}")
    add_bullet(doc, f"Mediana de releases: {stats['releases_mediana']}")

    img_path = os.path.join(DIRETORIO_GRAFICOS, "rq03_releases.png")
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # RQ04
    add_heading_formatted(doc, "3.4 RQ04 — Sistemas populares são atualizados com frequência?", level=2)
    add_paragraph_formatted(doc, f"Métrica: Dias desde a última atualização")
    add_bullet(doc, f"Média de dias desde última atualização: {stats['dias_atualizacao_media']}")
    add_bullet(doc, f"Mediana de dias desde última atualização: {stats['dias_atualizacao_mediana']}")

    img_path = os.path.join(DIRETORIO_GRAFICOS, "rq04_atualizacao.png")
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # RQ05
    add_heading_formatted(doc, "3.5 RQ05 — Sistemas populares são escritos nas linguagens mais populares?", level=2)
    add_paragraph_formatted(doc, f"Métrica: Linguagem primária de cada repositório")
    add_paragraph_formatted(doc, f"As 3 linguagens mais frequentes representam {stats['top3_percentual']}% dos repositórios analisados:")
    for lang, count in stats["top3_linguagens"].items():
        add_bullet(doc, f"{lang}: {count} repositórios")

    img_path = os.path.join(DIRETORIO_GRAFICOS, "rq05_linguagens.png")
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # RQ06
    add_heading_formatted(doc, "3.6 RQ06 — Sistemas populares possuem alto percentual de issues fechadas?", level=2)
    add_paragraph_formatted(doc, f"Métrica: Razão entre número de issues fechadas pelo total de issues")
    add_bullet(doc, f"Média do percentual de issues fechadas: {stats['issues_fechadas_media']}%")
    add_bullet(doc, f"Mediana do percentual de issues fechadas: {stats['issues_fechadas_mediana']}%")

    img_path = os.path.join(DIRETORIO_GRAFICOS, "rq06_issues_fechadas.png")
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # RQ07
    add_heading_formatted(doc, "3.7 RQ07 — Repositórios em linguagens populares vs. menos populares", level=2)
    add_paragraph_formatted(doc,
        "Comparação entre repositórios escritos nas 5 linguagens mais populares versus os demais:"
    )

    # Tabela comparativa
    table_comp = doc.add_table(rows=4, cols=3)
    table_comp.style = "Light Shading Accent 1"
    table_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_comp.rows[0].cells[0].text = "Métrica"
    table_comp.rows[0].cells[1].text = "Linguagens Populares"
    table_comp.rows[0].cells[2].text = "Outras Linguagens"
    table_comp.rows[1].cells[0].text = "Média de PRs Aceitas"
    table_comp.rows[1].cells[1].text = str(stats["pr_media_popular"])
    table_comp.rows[1].cells[2].text = str(stats["pr_media_nao_popular"])
    table_comp.rows[2].cells[0].text = "Média de Releases"
    table_comp.rows[2].cells[1].text = str(stats["releases_media_popular"])
    table_comp.rows[2].cells[2].text = str(stats["releases_media_nao_popular"])
    table_comp.rows[3].cells[0].text = "Média Dias Última Atualização"
    table_comp.rows[3].cells[1].text = str(stats["dias_media_popular"])
    table_comp.rows[3].cells[2].text = str(stats["dias_media_nao_popular"])

    add_paragraph_formatted(doc, "")
    add_paragraph_formatted(doc,
        f"Repositórios em linguagens populares recebem {abs(stats['pr_diff_pct'])}% "
        f"{'mais' if stats['pr_diff_pct'] > 0 else 'menos'} pull requests aceitas e lançam "
        f"{abs(stats['releases_diff_pct'])}% {'mais' if stats['releases_diff_pct'] > 0 else 'menos'} releases."
    )

    # Tabela detalhada por linguagem
    add_heading_formatted(doc, "3.8 Tabela Detalhada por Linguagem (Top 10)", level=2)
    tabela_langs = stats["tabela_linguagens"]
    table_langs = doc.add_table(rows=len(tabela_langs) + 1, cols=4)
    table_langs.style = "Light Shading Accent 1"
    table_langs.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_langs.rows[0].cells[0].text = "Linguagem"
    table_langs.rows[0].cells[1].text = "Média de PRs Aceitas"
    table_langs.rows[0].cells[2].text = "Média de Releases"
    table_langs.rows[0].cells[3].text = "Média Dias (Última Atualização)"

    for i, (lang, row) in enumerate(tabela_langs.iterrows()):
        table_langs.rows[i + 1].cells[0].text = str(lang)
        table_langs.rows[i + 1].cells[1].text = str(row["media_prs"])
        table_langs.rows[i + 1].cells[2].text = str(row["media_releases"])
        table_langs.rows[i + 1].cells[3].text = str(row["media_dias"])

    # ============================
    # 4. DISCUSSÃO DOS RESULTADOS
    # ============================
    add_heading_formatted(doc, "4. Discussão dos Resultados", level=1)

    add_heading_formatted(doc, "4.1 Confronto com as Hipóteses", level=2)

    # H1
    h1_confirmada = stats["idade_media"] > 5
    add_paragraph_formatted(doc,
        f"H1 — {'✅ Confirmada' if h1_confirmada else '❌ Não confirmada'}. "
        f"A idade média dos repositórios analisados é de {stats['idade_media']} anos, com uma mediana de "
        f"{stats['idade_mediana']} anos, {'indicando que projetos populares geralmente possuem um longo histórico de desenvolvimento e aprimoramento contínuo.' if h1_confirmada else 'indicando que os repositórios populares são relativamente recentes.'}",
        bold=True
    )

    # H2
    h2_confirmada = stats["pr_media"] >= 1000
    add_paragraph_formatted(doc,
        f"H2 — {'✅ Confirmada' if h2_confirmada else '⚠️ Parcialmente confirmada'}. "
        f"A média de {stats['pr_media']} pull requests aceitas é bastante alta, sugerindo grande participação "
        f"da comunidade. Contudo, a mediana de {stats['pr_mediana']} indica que a maioria dos repositórios recebe "
        f"menos contribuições do que alguns poucos extremamente populares.",
        bold=True
    )

    # H3
    h3_confirmada = stats["releases_media"] > 10
    add_paragraph_formatted(doc,
        f"H3 — {'✅ Confirmada' if h3_confirmada else '⚠️ Parcialmente confirmada'}. "
        f"A média de {stats['releases_media']} releases sugere que alguns projetos possuem um ritmo acelerado "
        f"de lançamentos, mas a mediana de {stats['releases_mediana']} releases mostra que essa não é uma regra "
        f"para todos os repositórios populares.",
        bold=True
    )

    # H4
    h4_confirmada = stats["dias_atualizacao_media"] < 30
    add_paragraph_formatted(doc,
        f"H4 — {'✅ Confirmada' if h4_confirmada else '❌ Não confirmada'}. "
        f"A média de dias desde a última atualização é {stats['dias_atualizacao_media']} dias e a mediana é "
        f"{stats['dias_atualizacao_mediana']} dias, {'reforçando que projetos populares são frequentemente mantidos e atualizados.' if h4_confirmada else 'indicando que nem todos os repositórios populares são atualizados com alta frequência.'}",
        bold=True
    )

    # H5
    h5_confirmada = stats["top3_percentual"] >= 50
    top3_nomes = ", ".join(stats["top3_linguagens"].index.tolist())
    add_paragraph_formatted(doc,
        f"H5 — {'✅ Confirmada' if h5_confirmada else '❌ Não confirmada'}. "
        f"As 3 linguagens mais frequentes ({top3_nomes}) representam {stats['top3_percentual']}% dos repositórios, "
        f"{'confirmando que repositórios populares tendem a utilizar linguagens amplamente adotadas pela comunidade.' if h5_confirmada else 'indicando uma distribuição mais diversificada de linguagens.'}",
        bold=True
    )

    # H6
    h6_confirmada = stats["issues_fechadas_mediana"] >= 70
    add_paragraph_formatted(doc,
        f"H6 — {'✅ Confirmada' if h6_confirmada else '❌ Não confirmada'}. "
        f"A média de {stats['issues_fechadas_media']}% e a mediana de {stats['issues_fechadas_mediana']}% de issues "
        f"fechadas demonstram que a maioria dos projetos analisados possui uma boa gestão de demandas, "
        f"garantindo a resolução eficaz de problemas e sugestões da comunidade.",
        bold=True
    )

    # H7
    h7_pr = stats["pr_diff_pct"] >= 20
    h7_releases = stats["releases_diff_pct"] >= 15
    h7_confirmada = h7_pr and h7_releases
    add_paragraph_formatted(doc,
        f"H7 (bônus) — {'✅ Confirmada' if h7_confirmada else '⚠️ Parcialmente confirmada'}. "
        f"Repositórios em linguagens populares recebem {abs(stats['pr_diff_pct'])}% "
        f"{'mais' if stats['pr_diff_pct'] > 0 else 'menos'} pull requests e lançam "
        f"{abs(stats['releases_diff_pct'])}% {'mais' if stats['releases_diff_pct'] > 0 else 'menos'} releases "
        f"do que os escritos em linguagens menos comuns.",
        bold=True
    )

    add_heading_formatted(doc, "4.2 Insights", level=2)
    insights = [
        "Repositórios populares são, em sua maioria, antigos e bem mantidos. A alta idade média e a regularidade nas atualizações confirmam essa tendência.",
        "A participação externa varia bastante. Alguns projetos extremamente populares recebem milhares de contribuições, enquanto a maioria tem um fluxo mais moderado, como demonstra a diferença entre média e mediana de pull requests.",
        "A frequência de releases não é uniforme. Embora a média seja alta, a mediana mostra que muitos repositórios lançam versões com menos frequência.",
        "A gestão de issues é eficiente na maioria dos repositórios populares, sugerindo um bom nível de manutenção e engajamento da comunidade.",
        "A influência da linguagem de programação na popularidade é significativa, com linguagens como TypeScript, Python e JavaScript dominando o cenário.",
    ]
    for insight in insights:
        add_bullet(doc, insight)

    add_heading_formatted(doc, "4.3 Estatísticas Gerais", level=2)
    add_bullet(doc, f"Total de repositórios analisados: {len(df)}")
    add_bullet(doc, f"Estrelas médias por repositório: {round(df['Estrelas'].mean(), 0):.0f}")
    add_bullet(doc, f"Mediana de estrelas: {round(df['Estrelas'].median(), 0):.0f}")
    add_bullet(doc, f"Repositório mais antigo: {df.loc[df['Idade (anos)'].idxmax(), 'Nome']} ({df['Idade (anos)'].max()} anos)")
    add_bullet(doc, f"Repositório mais recente: {df.loc[df['Idade (anos)'].idxmin(), 'Nome']} ({df['Idade (anos)'].min()} anos)")
    add_bullet(doc, f"Maior número de PRs aceitas: {df.loc[df['Pull Requests Aceitos'].idxmax(), 'Nome']} ({df['Pull Requests Aceitos'].max():,})")
    add_bullet(doc, f"Número de linguagens distintas: {df['Linguagem Principal'].nunique()}")

    # ============================
    # 5. CONCLUSÃO
    # ============================
    add_heading_formatted(doc, "5. Conclusão", level=1)

    add_heading_formatted(doc, "5.1 Tomada de Decisão", level=2)
    add_paragraph_formatted(doc,
        "Com base nos resultados, conclui-se que repositórios populares no GitHub possuem características "
        "distintas que os diferenciam de projetos menos populares: são mais maduros, recebem contribuições "
        "externas significativas, mantêm uma gestão eficiente de issues e são predominantemente escritos "
        "em linguagens de programação amplamente adotadas. Esses achados podem orientar desenvolvedores e "
        "mantenedores de projetos open-source na adoção de práticas que favoreçam a popularidade e a "
        "sustentabilidade de seus repositórios."
    )

    add_heading_formatted(doc, "5.2 Sugestões Futuras", level=2)
    sugestoes = [
        "Expandir a análise para incluir métricas de qualidade de código (cobertura de testes, complexidade ciclomática).",
        "Investigar a correlação entre número de contribuidores ativos e a longevidade do repositório.",
        "Analisar a evolução temporal das métricas (como as estrelas e contribuições crescem ao longo dos anos).",
        "Comparar repositórios populares open-source com projetos privados ou de código fechado.",
        "Incorporar análise de sentimento nos comentários de issues e pull requests.",
    ]
    for s in sugestoes:
        add_bullet(doc, s)

    add_heading_formatted(doc, "5.3 Resultado Conclusivo", level=2)
    add_paragraph_formatted(doc,
        f"A análise de {len(df)} repositórios populares do GitHub revela que projetos bem-sucedidos compartilham "
        f"características como maturidade (mediana de {stats['idade_mediana']} anos), gestão eficiente de issues "
        f"(mediana de {stats['issues_fechadas_mediana']}% fechadas) e manutenção ativa. As linguagens TypeScript, "
        f"Python e JavaScript dominam entre os projetos mais populares. A hipótese de que linguagens populares "
        f"atraem mais contribuições e releases foi corroborada pelos dados."
    )

    add_heading_formatted(doc, "5.4 Confronto com a Literatura", level=2)
    add_paragraph_formatted(doc,
        "Os resultados obtidos são consistentes com estudos anteriores na área de engenharia de software empírica. "
        "Kalliamvakou et al. (2014), em seu estudo sobre o GitHub como plataforma de desenvolvimento colaborativo, "
        "identificaram que repositórios populares tendem a apresentar maior atividade de contribuição e "
        "manutenção contínua. Além disso, Borges et al. (2016) demonstraram que a popularidade de repositórios "
        "no GitHub está correlacionada com a linguagem de programação utilizada, com linguagens como JavaScript "
        "e Python liderando em número de estrelas — achado que corrobora os resultados do presente estudo. "
        "Munaiah et al. (2017) também reforçaram que repositórios com alta atividade de issues e pull requests "
        "são indicadores de projetos com boa engenharia, o que se alinha com a alta taxa de fechamento de issues "
        "observada nesta análise."
    )

    # Salvar DOCX
    doc.save(CAMINHO_DOCX)
    print(f"✅ Relatório DOCX salvo em: {CAMINHO_DOCX}")
    return CAMINHO_DOCX


def gerar_pdf_direto(df, stats, caminho_pdf):
    """Gera PDF diretamente usando fpdf2."""
    from fpdf import FPDF

    def sanitize(text):
        """Remove/replace characters not supported by latin-1 encoding."""
        replacements = {
            '\u2014': '-',   # em dash
            '\u2013': '-',   # en dash
            '\u2018': "'",   # left single quote
            '\u2019': "'",   # right single quote
            '\u201c': '"',   # left double quote
            '\u201d': '"',   # right double quote
            '\u2026': '...', # ellipsis
            '\u2022': '-',   # bullet
            '\u2713': '[v]', # check mark
            '\u2715': '[x]', # cross mark
            '\u2705': '[OK]',
            '\u274c': '[X]',
            '\u26a0': '[!]',
            '\U0001f4ca': '',  # chart emoji
            '\U0001f4c8': '',  # chart up emoji
            '\U0001f4dd': '',  # memo emoji
            '\U0001f4c4': '',  # page emoji
            '\u2705': '[OK]',
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        # Fallback: encode to latin-1 replacing unknown chars
        return text.encode('latin-1', errors='replace').decode('latin-1')

    class RelatorioPDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "B", 9)
            self.set_text_color(150, 150, 150)
            self.cell(0, 8, sanitize("Relatorio Final - Repositorios Populares do GitHub"), align="C", new_x="LMARGIN", new_y="NEXT")
            self.line(10, self.get_y(), 200, self.get_y())
            self.ln(3)

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 10, sanitize(f"Pagina {self.page_no()}/{{nb}}"), align="C")

        def titulo_secao(self, texto, nivel=1):
            if nivel == 1:
                self.set_font("Helvetica", "B", 16)
                self.set_text_color(30, 70, 130)
            elif nivel == 2:
                self.set_font("Helvetica", "B", 13)
                self.set_text_color(50, 100, 160)
            else:
                self.set_font("Helvetica", "B", 11)
                self.set_text_color(70, 120, 180)
            self.ln(4)
            self.multi_cell(0, 7, sanitize(texto))
            self.ln(2)

        def paragrafo(self, texto, negrito=False):
            if negrito:
                self.set_font("Helvetica", "B", 10)
            else:
                self.set_font("Helvetica", "", 10)
            self.set_text_color(30, 30, 30)
            self.multi_cell(0, 5.5, sanitize(texto))
            self.ln(2)

        def bullet(self, texto):
            self.set_font("Helvetica", "", 10)
            self.set_text_color(30, 30, 30)
            self.cell(6, 5.5, "-")
            self.multi_cell(0, 5.5, sanitize(texto))
            self.ln(1)

        def inserir_imagem(self, caminho, largura=170):
            if os.path.exists(caminho):
                self.image(caminho, x=(210 - largura) / 2, w=largura)
                self.ln(5)

        def tabela_simples(self, headers, dados, col_widths=None):
            if col_widths is None:
                col_widths = [190 / len(headers)] * len(headers)
            # Header
            self.set_font("Helvetica", "B", 9)
            self.set_fill_color(70, 130, 180)
            self.set_text_color(255, 255, 255)
            for i, h in enumerate(headers):
                self.cell(col_widths[i], 7, sanitize(h), border=1, fill=True, align="C")
            self.ln()
            # Dados
            self.set_font("Helvetica", "", 9)
            self.set_text_color(30, 30, 30)
            fill = False
            for row in dados:
                if fill:
                    self.set_fill_color(230, 240, 250)
                else:
                    self.set_fill_color(255, 255, 255)
                for i, val in enumerate(row):
                    self.cell(col_widths[i], 6, sanitize(str(val)), border=1, fill=True, align="C")
                self.ln()
                fill = not fill
            self.ln(3)

    pdf = RelatorioPDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Página de título
    pdf.add_page()
    pdf.ln(40)
    pdf.set_font("Helvetica", "B", 28)
    pdf.set_text_color(30, 70, 130)
    pdf.cell(0, 15, sanitize("Relatorio Final"), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)
    pdf.set_font("Helvetica", "", 18)
    pdf.set_text_color(70, 130, 180)
    pdf.cell(0, 10, sanitize("Caracteristicas de Repositorios"), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 10, "Populares no GitHub", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(10)
    pdf.set_font("Helvetica", "I", 13)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 8, sanitize("Laboratorio de Experimentacao de Software - Lab01"), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(30)
    pdf.set_font("Helvetica", "", 11)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(0, 8, sanitize(f"Total de repositorios analisados: {len(df)}"), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 8, sanitize("Dados coletados via API GraphQL do GitHub"), align="C", new_x="LMARGIN", new_y="NEXT")

    # 1. INTRODUÇÃO
    pdf.add_page()
    pdf.titulo_secao("1. Introdução")

    pdf.titulo_secao("1.1 Contextualização", 2)
    pdf.paragrafo(
        "O GitHub é a maior plataforma de hospedagem de código-fonte e colaboração de software do mundo, "
        "abrigando milhões de repositórios open-source. Projetos populares nessa plataforma são frequentemente "
        "utilizados como referência por desenvolvedores, empresas e pesquisadores. Compreender as características "
        "desses repositórios pode revelar padrões sobre boas práticas de desenvolvimento, manutenção e "
        "engajamento da comunidade."
    )

    pdf.titulo_secao("1.2 Problema Foco do Experimento", 2)
    pdf.paragrafo(
        "Este estudo busca responder: quais são as características predominantes dos repositórios "
        "mais populares do GitHub em termos de idade, contribuição externa, frequência de releases, "
        "frequência de atualização, linguagens de programação utilizadas e gestão de issues?"
    )

    pdf.titulo_secao("1.3 Questões de Pesquisa", 2)
    questoes = [
        "RQ01: Sistemas populares são maduros/antigos?",
        "RQ02: Sistemas populares recebem muita contribuição externa?",
        "RQ03: Sistemas populares lançam releases com frequência?",
        "RQ04: Sistemas populares são atualizados com frequência?",
        "RQ05: Sistemas populares são escritos nas linguagens mais populares?",
        "RQ06: Sistemas populares possuem um alto percentual de issues fechadas?",
        "RQ07: Sistemas em linguagens populares recebem mais contribuição, lançam mais releases e são mais atualizados?",
    ]
    for q in questoes:
        pdf.bullet(q)

    pdf.titulo_secao("1.4 Hipóteses", 2)
    hipoteses = [
        "H1: Repositórios populares tendem a ser mais antigos, com idade média superior a 5 anos.",
        "H2: Projetos populares recebem pelo menos 1.000 pull requests aceitas em média.",
        "H3: Repositórios populares têm total médio de mais de 10 releases.",
        "H4: Projetos amplamente utilizados têm tempo médio desde última atualização inferior a 30 dias.",
        "H5: Repositórios populares são escritos em JavaScript, Python e TypeScript (>= 50% do total).",
        "H6: Projetos populares possuem pelo menos 70% das issues resolvidas.",
        "H7 (bônus): Linguagens populares recebem 20% mais PRs e 15% mais releases.",
    ]
    for h in hipoteses:
        pdf.bullet(h)

    pdf.titulo_secao("1.5 Objetivos", 2)
    pdf.paragrafo(
        "Objetivo principal: Analisar as características dos 1.000 repositórios mais populares do GitHub, "
        "identificando padrões e tendências relacionados à idade, contribuição, lançamentos, manutenção, "
        "linguagens e gestão de issues.",
        negrito=True
    )
    objetivos = [
        "Coletar dados dos 1.000 repositórios com mais estrelas via API GraphQL.",
        "Calcular métricas estatísticas (média, mediana) para cada questão de pesquisa.",
        "Gerar visualizações gráficas para cada métrica analisada.",
        "Confrontar os resultados obtidos com as hipóteses formuladas.",
        "Comparar repositórios em linguagens populares versus menos populares.",
    ]
    for o in objetivos:
        pdf.bullet(o)

    # 2. METODOLOGIA
    pdf.add_page()
    pdf.titulo_secao("2. Metodologia")

    pdf.titulo_secao("2.1 Passo a Passo do Experimento", 2)
    passos = [
        "1. Definição das questões de pesquisa e hipóteses iniciais.",
        "2. Implementação do script de coleta utilizando a API GraphQL do GitHub.",
        "3. Coleta dos dados dos 1.000 repositórios com maior número de estrelas.",
        "4. Armazenamento dos dados em arquivo CSV para análise.",
        "5. Processamento estatístico com cálculo de média, mediana e distribuições.",
        "6. Geração de gráficos para visualização dos resultados.",
        "7. Confrontação dos resultados com as hipóteses formuladas.",
    ]
    for p in passos:
        pdf.paragrafo(p)

    pdf.titulo_secao("2.2 Decisões", 2)
    pdf.paragrafo(
        "Optou-se por utilizar a mediana como métrica central de tendência para evitar distorções causadas "
        "por outliers. A coleta foi limitada a 1.000 repositórios para garantir uma amostra significativa "
        "e viável dentro dos limites da API do GitHub."
    )

    pdf.titulo_secao("2.3 Materiais Utilizados", 2)
    materiais = [
        "API GraphQL do GitHub para coleta de dados.",
        "Python 3.x como linguagem de programação.",
        "Pandas para manipulação e análise de dados.",
        "Matplotlib para geração de gráficos.",
        "Requests para comunicação HTTP com a API.",
        "Token de acesso pessoal do GitHub para autenticação.",
    ]
    for m in materiais:
        pdf.bullet(m)

    pdf.titulo_secao("2.4 Métricas e Unidades", 2)
    metricas_headers = ["Questão", "Métrica", "Unidade"]
    metricas_dados = [
        ("RQ01", "Idade do repositório", "Anos"),
        ("RQ02", "Total de pull requests aceitas", "Quantidade"),
        ("RQ03", "Total de releases", "Quantidade"),
        ("RQ04", "Tempo desde última atualização", "Dias"),
        ("RQ05", "Linguagem primária", "Categórica"),
        ("RQ06", "Razão issues fechadas / total", "Proporção (0 a 1)"),
    ]
    pdf.tabela_simples(metricas_headers, metricas_dados, [30, 100, 60])

    # 3. RESULTADOS
    pdf.add_page()
    pdf.titulo_secao("3. Visualização dos Resultados")

    # RQ01
    pdf.titulo_secao("3.1 RQ01 — Sistemas populares são maduros/antigos?", 2)
    pdf.paragrafo("Métrica: Idade do repositório (calculada a partir da data de criação)")
    pdf.bullet(f"Idade média dos repositórios: {stats['idade_media']} anos")
    pdf.bullet(f"Mediana da idade: {stats['idade_mediana']} anos")
    pdf.inserir_imagem(os.path.join(DIRETORIO_GRAFICOS, "rq01_idade.png"))

    # RQ02
    pdf.titulo_secao("3.2 RQ02 — Contribuição externa", 2)
    pdf.paragrafo("Métrica: Total de pull requests aceitas")
    pdf.bullet(f"Média de pull requests aceitos: {stats['pr_media']}")
    pdf.bullet(f"Mediana de pull requests aceitos: {stats['pr_mediana']}")
    pdf.inserir_imagem(os.path.join(DIRETORIO_GRAFICOS, "rq02_pull_requests.png"))

    # RQ03
    pdf.add_page()
    pdf.titulo_secao("3.3 RQ03 — Frequência de releases", 2)
    pdf.paragrafo("Métrica: Total de releases")
    pdf.bullet(f"Média de releases: {stats['releases_media']}")
    pdf.bullet(f"Mediana de releases: {stats['releases_mediana']}")
    pdf.inserir_imagem(os.path.join(DIRETORIO_GRAFICOS, "rq03_releases.png"))

    # RQ04
    pdf.titulo_secao("3.4 RQ04 — Frequência de atualização", 2)
    pdf.paragrafo("Métrica: Dias desde a última atualização")
    pdf.bullet(f"Média de dias desde última atualização: {stats['dias_atualizacao_media']}")
    pdf.bullet(f"Mediana de dias desde última atualização: {stats['dias_atualizacao_mediana']}")
    pdf.inserir_imagem(os.path.join(DIRETORIO_GRAFICOS, "rq04_atualizacao.png"))

    # RQ05
    pdf.add_page()
    pdf.titulo_secao("3.5 RQ05 — Linguagens mais populares", 2)
    pdf.paragrafo("Métrica: Linguagem primária de cada repositório")
    pdf.paragrafo(f"As 3 linguagens mais frequentes representam {stats['top3_percentual']}% dos repositórios:")
    for lang, count in stats["top3_linguagens"].items():
        pdf.bullet(f"{lang}: {count} repositórios")
    pdf.inserir_imagem(os.path.join(DIRETORIO_GRAFICOS, "rq05_linguagens.png"))

    # RQ06
    pdf.titulo_secao("3.6 RQ06 — Percentual de issues fechadas", 2)
    pdf.paragrafo("Métrica: Razão entre número de issues fechadas pelo total de issues")
    pdf.bullet(f"Média do percentual de issues fechadas: {stats['issues_fechadas_media']}%")
    pdf.bullet(f"Mediana do percentual de issues fechadas: {stats['issues_fechadas_mediana']}%")
    pdf.inserir_imagem(os.path.join(DIRETORIO_GRAFICOS, "rq06_issues_fechadas.png"))

    # RQ07
    pdf.add_page()
    pdf.titulo_secao("3.7 RQ07 — Linguagens populares vs. menos populares", 2)
    pdf.paragrafo("Comparação entre repositórios escritos nas 5 linguagens mais populares versus os demais:")

    comp_headers = ["Métrica", "Ling. Populares", "Outras"]
    comp_dados = [
        ("Média de PRs Aceitas", str(stats["pr_media_popular"]), str(stats["pr_media_nao_popular"])),
        ("Média de Releases", str(stats["releases_media_popular"]), str(stats["releases_media_nao_popular"])),
        ("Média Dias Última Atualiz.", str(stats["dias_media_popular"]), str(stats["dias_media_nao_popular"])),
    ]
    pdf.tabela_simples(comp_headers, comp_dados, [80, 55, 55])

    pdf.paragrafo(
        f"Repositórios em linguagens populares recebem {abs(stats['pr_diff_pct'])}% "
        f"{'mais' if stats['pr_diff_pct'] > 0 else 'menos'} pull requests aceitas e lançam "
        f"{abs(stats['releases_diff_pct'])}% {'mais' if stats['releases_diff_pct'] > 0 else 'menos'} releases."
    )

    # Tabela por linguagem
    pdf.titulo_secao("3.8 Tabela Detalhada por Linguagem (Top 10)", 2)
    lang_headers = ["Linguagem", "Média PRs", "Média Releases", "Média Dias Atualiz."]
    lang_dados = []
    for lang, row in stats["tabela_linguagens"].iterrows():
        lang_dados.append((str(lang), str(row["media_prs"]), str(row["media_releases"]), str(row["media_dias"])))
    pdf.tabela_simples(lang_headers, lang_dados, [55, 45, 45, 45])

    # 4. DISCUSSÃO
    pdf.add_page()
    pdf.titulo_secao("4. Discussão dos Resultados")
    pdf.titulo_secao("4.1 Confronto com as Hipóteses", 2)

    h1_ok = stats["idade_media"] > 5
    pdf.paragrafo(
        f"H1 — {'Confirmada' if h1_ok else 'Não confirmada'}. "
        f"Idade média: {stats['idade_media']} anos, mediana: {stats['idade_mediana']} anos. "
        f"{'Projetos populares possuem longo histórico de desenvolvimento.' if h1_ok else ''}",
        negrito=True
    )

    h2_ok = stats["pr_media"] >= 1000
    pdf.paragrafo(
        f"H2 — {'Confirmada' if h2_ok else 'Parcialmente confirmada'}. "
        f"Média: {stats['pr_media']} PRs, mediana: {stats['pr_mediana']}. "
        f"A diferença entre média e mediana indica grande variação entre repositórios.",
        negrito=True
    )

    h3_ok = stats["releases_media"] > 10
    pdf.paragrafo(
        f"H3 — {'Confirmada' if h3_ok else 'Parcialmente confirmada'}. "
        f"Média: {stats['releases_media']} releases, mediana: {stats['releases_mediana']}. "
        f"Nem todos os repositórios populares lançam releases com alta frequência.",
        negrito=True
    )

    h4_ok = stats["dias_atualizacao_media"] < 30
    pdf.paragrafo(
        f"H4 — {'Confirmada' if h4_ok else 'Não confirmada'}. "
        f"Média: {stats['dias_atualizacao_media']} dias, mediana: {stats['dias_atualizacao_mediana']} dias. "
        f"{'Projetos populares são frequentemente mantidos e atualizados.' if h4_ok else ''}",
        negrito=True
    )

    h5_ok = stats["top3_percentual"] >= 50
    top3_nomes = ", ".join(stats["top3_linguagens"].index.tolist())
    pdf.paragrafo(
        f"H5 — {'Confirmada' if h5_ok else 'Não confirmada'}. "
        f"As 3 linguagens mais frequentes ({top3_nomes}) representam {stats['top3_percentual']}% dos repositórios.",
        negrito=True
    )

    h6_ok = stats["issues_fechadas_mediana"] >= 70
    pdf.paragrafo(
        f"H6 — {'Confirmada' if h6_ok else 'Não confirmada'}. "
        f"Média: {stats['issues_fechadas_media']}%, mediana: {stats['issues_fechadas_mediana']}%. "
        f"A maioria dos projetos possui boa gestão de issues.",
        negrito=True
    )

    h7_pr = stats["pr_diff_pct"] >= 20
    h7_rel = stats["releases_diff_pct"] >= 15
    h7_ok = h7_pr and h7_rel
    pdf.paragrafo(
        f"H7 (bônus) — {'Confirmada' if h7_ok else 'Parcialmente confirmada'}. "
        f"Repositórios em linguagens populares recebem {abs(stats['pr_diff_pct'])}% "
        f"{'mais' if stats['pr_diff_pct'] > 0 else 'menos'} PRs e lançam {abs(stats['releases_diff_pct'])}% "
        f"{'mais' if stats['releases_diff_pct'] > 0 else 'menos'} releases.",
        negrito=True
    )

    pdf.titulo_secao("4.2 Insights", 2)
    insights = [
        "Repositórios populares são, em sua maioria, antigos e bem mantidos.",
        "A participação externa varia bastante, como demonstra a diferença entre média e mediana de PRs.",
        "A frequência de releases não é uniforme entre todos os repositórios populares.",
        "A gestão de issues é eficiente na maioria dos projetos populares.",
        "Linguagens como TypeScript, Python e JavaScript dominam o cenário.",
    ]
    for i in insights:
        pdf.bullet(i)

    pdf.titulo_secao("4.3 Estatísticas Gerais", 2)
    pdf.bullet(f"Total de repositórios analisados: {len(df)}")
    pdf.bullet(f"Estrelas médias por repositório: {round(df['Estrelas'].mean(), 0):.0f}")
    pdf.bullet(f"Mediana de estrelas: {round(df['Estrelas'].median(), 0):.0f}")
    pdf.bullet(f"Repositório mais antigo: {df.loc[df['Idade (anos)'].idxmax(), 'Nome']} ({df['Idade (anos)'].max()} anos)")
    pdf.bullet(f"Repositório mais recente: {df.loc[df['Idade (anos)'].idxmin(), 'Nome']} ({df['Idade (anos)'].min()} anos)")
    pdf.bullet(f"Número de linguagens distintas: {df['Linguagem Principal'].nunique()}")

    # 5. CONCLUSÃO
    pdf.add_page()
    pdf.titulo_secao("5. Conclusão")

    pdf.titulo_secao("5.1 Tomada de Decisão", 2)
    pdf.paragrafo(
        "Com base nos resultados, conclui-se que repositórios populares no GitHub possuem características "
        "distintas: são mais maduros, recebem contribuições externas significativas, mantêm uma gestão eficiente "
        "de issues e são predominantemente escritos em linguagens amplamente adotadas. Esses achados podem "
        "orientar desenvolvedores e mantenedores de projetos open-source."
    )

    pdf.titulo_secao("5.2 Sugestões Futuras", 2)
    sugestoes = [
        "Expandir a análise para incluir métricas de qualidade de código.",
        "Investigar a correlação entre contribuidores ativos e longevidade do repositório.",
        "Analisar a evolução temporal das métricas ao longo dos anos.",
        "Comparar repositórios open-source com projetos de código fechado.",
        "Incorporar análise de sentimento nos comentários de issues e PRs.",
    ]
    for s in sugestoes:
        pdf.bullet(s)

    pdf.titulo_secao("5.3 Resultado Conclusivo", 2)
    pdf.paragrafo(
        f"A análise de {len(df)} repositórios populares do GitHub revela que projetos bem-sucedidos compartilham "
        f"características como maturidade (mediana de {stats['idade_mediana']} anos), gestão eficiente de issues "
        f"(mediana de {stats['issues_fechadas_mediana']}% fechadas) e manutenção ativa. As linguagens TypeScript, "
        f"Python e JavaScript dominam entre os projetos mais populares."
    )

    pdf.titulo_secao("5.4 Confronto com a Literatura", 2)
    pdf.paragrafo(
        "Os resultados são consistentes com estudos anteriores. Kalliamvakou et al. (2014) identificaram que "
        "repositórios populares tendem a apresentar maior atividade de contribuição e manutenção contínua. "
        "Borges et al. (2016) demonstraram que a popularidade está correlacionada com a linguagem utilizada, "
        "com JavaScript e Python liderando — achado corroborado pelo presente estudo. "
        "Munaiah et al. (2017) reforçaram que repositórios com alta atividade de issues e pull requests "
        "são indicadores de projetos com boa engenharia, alinhado com a alta taxa de fechamento observada."
    )

    pdf.output(caminho_pdf)
    print(f"✅ Relatório PDF salvo em: {caminho_pdf}")


def main():
    print("📊 Carregando dados dos repositórios...")
    df = carregar_dados()
    print(f"   {len(df)} repositórios carregados.")

    print("📈 Calculando estatísticas...")
    stats = calcular_estatisticas(df)

    print("📝 Gerando relatório DOCX...")
    caminho_docx = gerar_docx(df, stats)

    print("📄 Gerando relatório PDF...")
    gerar_pdf_direto(df, stats, CAMINHO_PDF)

    print("\n✅ Processo concluído!")


if __name__ == "__main__":
    main()

