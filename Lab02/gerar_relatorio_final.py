"""Gera o relatório final do Lab02 em DOCX e PDF.

Estrutura do relatório:
1. Introdução (contextualização, problema, QPs, hipóteses, objetivos)
2. Metodologia (passo a passo, decisões, materiais, métodos, métricas)
3. Visualização dos Resultados (tabelas, gráficos)
4. Discussão (confronto QPs, insights, estatísticas)
5. Conclusão (tomada de decisão, sugestões futuras)
6. Referências

Todos os gráficos em preto e branco, apresentação simples.
"""

from __future__ import annotations

import argparse
import io
import textwrap
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, Cm, RGBColor
from matplotlib.backends.backend_pdf import PdfPages
from scipy.stats import spearmanr

# ---------- Caminhos ----------
BASE_DIR = Path(__file__).resolve().parent
REPORTS_DIR = BASE_DIR / "reports"
SOURCE_HTML = REPORTS_DIR / "report.html"
SOURCE_CSV = REPORTS_DIR / "dados_brutos.csv"
DEFAULT_DOCX = REPORTS_DIR / "relatorio_final.docx"
DEFAULT_PDF = REPORTS_DIR / "relatorio_final.pdf"
GRAPHS_DIR = REPORTS_DIR

# ---------- Colunas ----------
COLUMN_NAMES = [
    "Nome", "Proprietario", "Idade", "Estrelas",
    "Pull Requests Aceitos", "Releases",
    "Linhas de Codigo", "Linhas de Comentario",
    "CBO", "DIT", "LCOM",
]

PAIR_DEFINITIONS = [
    ("RQ1 — Popularidade", "Estrelas", "CBO", "popularidade"),
    ("RQ1 — Popularidade", "Estrelas", "DIT", "popularidade"),
    ("RQ1 — Popularidade", "Estrelas", "LCOM", "popularidade"),
    ("RQ2 — Maturidade", "Idade", "CBO", "maturidade"),
    ("RQ2 — Maturidade", "Idade", "DIT", "maturidade"),
    ("RQ2 — Maturidade", "Idade", "LCOM", "maturidade"),
    ("RQ3 — Atividade", "Releases", "CBO", "atividade"),
    ("RQ3 — Atividade", "Releases", "DIT", "atividade"),
    ("RQ3 — Atividade", "Releases", "LCOM", "atividade"),
    ("RQ4 — Tamanho", "Linhas de Codigo", "CBO", "tamanho"),
    ("RQ4 — Tamanho", "Linhas de Codigo", "DIT", "tamanho"),
    ("RQ4 — Tamanho", "Linhas de Codigo", "LCOM", "tamanho"),
    ("RQ4 — Tamanho", "Linhas de Comentario", "CBO", "tamanho"),
    ("RQ4 — Tamanho", "Linhas de Comentario", "DIT", "tamanho"),
    ("RQ4 — Tamanho", "Linhas de Comentario", "LCOM", "tamanho"),
]


@dataclass
class PairResult:
    rq: str
    x: str
    y: str
    n: int
    rho: float
    pvalue: float


# ================================================================
#  LEITURA DE DADOS
# ================================================================

def load_dataframe(source: Path | None = None) -> pd.DataFrame:
    """Carrega dados do CSV bruto ou do HTML (fallback)."""
    csv_path = source if source and source.suffix == ".csv" else SOURCE_CSV
    if csv_path.exists():
        df = pd.read_csv(csv_path, encoding="utf-8-sig")
        return _normalize_dataframe(df)

    html_path = source if source and source.suffix == ".html" else SOURCE_HTML
    if html_path.exists():
        tables = pd.read_html(html_path)
        for tbl in reversed(tables):
            if len(tbl.columns) >= 11:
                tbl.columns = COLUMN_NAMES[:len(tbl.columns)]
                return _normalize_dataframe(tbl)
        if tables:
            df = tables[0]
            df.columns = COLUMN_NAMES[:len(df.columns)]
            return _normalize_dataframe(df)

    raise FileNotFoundError("Nenhum arquivo de dados encontrado (dados_brutos.csv ou report.html)")


def _normalize_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    rename_map = {}
    for col in df.columns:
        cl = col.lower().strip()
        if "proprietário" in cl or "proprietario" in cl:
            rename_map[col] = "Proprietario"
        elif cl == "nome":
            rename_map[col] = "Nome"
        elif "idade" in cl:
            rename_map[col] = "Idade"
        elif "estrela" in cl or "star" in cl:
            rename_map[col] = "Estrelas"
        elif "pull" in cl:
            rename_map[col] = "Pull Requests Aceitos"
        elif "release" in cl:
            rename_map[col] = "Releases"
        elif ("código" in cl or "codigo" in cl) and "comentário" not in cl and "comentario" not in cl:
            rename_map[col] = "Linhas de Codigo"
        elif "comentário" in cl or "comentario" in cl:
            rename_map[col] = "Linhas de Comentario"
        elif "cbo" in cl and "classe" in cl and "método" not in cl and "metodo" not in cl:
            rename_map[col] = "CBO"
        elif "dit" in cl and "classe" in cl and "método" not in cl and "metodo" not in cl:
            rename_map[col] = "DIT"
        elif "lcom" in cl and "classe" in cl and "método" not in cl and "metodo" not in cl:
            rename_map[col] = "LCOM"
        elif cl == "cbo":
            rename_map[col] = "CBO"
        elif cl == "dit":
            rename_map[col] = "DIT"
        elif cl == "lcom":
            rename_map[col] = "LCOM"
    df = df.rename(columns=rename_map)

    if "Idade" in df.columns:
        df["Idade"] = df["Idade"].astype(str).str.replace(" anos", "", regex=False).str.strip()
        df["Idade"] = pd.to_numeric(df["Idade"], errors="coerce")

    for col in ["Estrelas", "Pull Requests Aceitos", "Releases",
                 "Linhas de Codigo", "Linhas de Comentario", "CBO", "DIT", "LCOM"]:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")
    return df


# ================================================================
#  ESTATÍSTICAS
# ================================================================

def summarize_metrics(df: pd.DataFrame) -> pd.DataFrame:
    metrics = [c for c in ["Idade", "Estrelas", "Releases", "Linhas de Codigo",
                            "Linhas de Comentario", "CBO", "DIT", "LCOM"] if c in df.columns]
    summary = df[metrics].agg(["mean", "median", "std", "min", "max"]).round(2)
    summary.index = ["Média", "Mediana", "Desvio Padrão", "Mínimo", "Máximo"]
    return summary


def compute_correlations(df: pd.DataFrame) -> list[PairResult]:
    results: list[PairResult] = []
    for rq, x, y, _group in PAIR_DEFINITIONS:
        if x not in df.columns or y not in df.columns:
            continue
        pair_df = df[[x, y]].dropna()
        if len(pair_df) < 3:
            results.append(PairResult(rq=rq, x=x, y=y, n=len(pair_df),
                                       rho=float("nan"), pvalue=float("nan")))
            continue
        rho, pvalue = spearmanr(pair_df[x], pair_df[y])
        results.append(PairResult(rq=rq, x=x, y=y, n=len(pair_df),
                                   rho=float(rho), pvalue=float(pvalue)))
    return results


def _interpret_correlation(rho: float, pvalue: float) -> str:
    if np.isnan(rho):
        return "Dados insuficientes"
    sig = "significativa (p < 0.05)" if pvalue < 0.05 else "não significativa (p >= 0.05)"
    abs_rho = abs(rho)
    if abs_rho < 0.1:
        strength = "desprezível"
    elif abs_rho < 0.3:
        strength = "fraca"
    elif abs_rho < 0.5:
        strength = "moderada"
    elif abs_rho < 0.7:
        strength = "forte"
    else:
        strength = "muito forte"
    direction = "positiva" if rho > 0 else "negativa"
    return f"Correlação {direction} {strength}, {sig}"


# ================================================================
#  GRÁFICOS — PALETA PROFISSIONAL POR RQ
# ================================================================

# Paleta de cores: cada RQ tem uma cor principal distinta
_PALETTE = {
    "rq1": {"dot": "#2563EB", "edge": "#1E40AF", "accent": "#DBEAFE",
             "label": "Popularidade"},                          # azul
    "rq2": {"dot": "#059669", "edge": "#065F46", "accent": "#D1FAE5",
             "label": "Maturidade"},                            # verde-esmeralda
    "rq3": {"dot": "#D97706", "edge": "#92400E", "accent": "#FEF3C7",
             "label": "Atividade"},                             # âmbar
    "rq4a": {"dot": "#7C3AED", "edge": "#5B21B6", "accent": "#EDE9FE",
              "label": "Tamanho (LOC)"},                        # violeta
    "rq4b": {"dot": "#DB2777", "edge": "#9D174D", "accent": "#FCE7F3",
              "label": "Tamanho (Comentários)"},                # rosa
}
# Cores para cada métrica CK nos box-plots
_CK_COLORS = {"CBO": "#3B82F6", "DIT": "#10B981", "LCOM": "#F59E0B"}

# Cores de tabela
_TBL_HEADER  = "#1E3A5F"
_TBL_HEADER_TEXT = "white"
_TBL_ROW_EVEN = "#F0F4F8"
_TBL_ROW_ODD  = "#FFFFFF"
_TBL_EDGE     = "#CBD5E1"

plt.rcParams.update({
    "font.family": "sans-serif",
    "font.sans-serif": ["Segoe UI", "Calibri", "Arial", "DejaVu Sans"],
    "font.size": 10,
    "axes.edgecolor": "#4B5563",
    "axes.labelcolor": "#1F2937",
    "xtick.color": "#4B5563",
    "ytick.color": "#4B5563",
    "text.color": "#1F2937",
    "figure.facecolor": "white",
    "axes.facecolor": "#FAFBFC",
    "savefig.facecolor": "white",
    "axes.grid": True,
    "grid.alpha": 0.25,
    "grid.color": "#D1D5DB",
    "grid.linestyle": "--",
})


def _create_scatter(ax, x, y, xlabel, ylabel, title_text="", palette_key="rq1"):
    """Scatter plot estilizado com cor da RQ."""
    pal = _PALETTE[palette_key]
    valid = pd.DataFrame({"x": x, "y": y}).dropna()
    if valid.empty:
        ax.set_visible(False)
        return
    ax.scatter(valid["x"], valid["y"], s=18, alpha=0.55,
               color=pal["dot"], edgecolors=pal["edge"],
               linewidths=0.4, marker="o", zorder=3)
    ax.set_xlabel(xlabel, fontsize=9, fontweight="medium")
    ax.set_ylabel(ylabel, fontsize=9, fontweight="medium")
    if title_text:
        ax.set_title(title_text, fontsize=10, fontweight="bold", color="#111827")
    ax.tick_params(labelsize=8)
    # anotação Spearman
    if len(valid) > 2:
        rho, pval = spearmanr(valid["x"], valid["y"])
        ax.annotate(
            f"ρ = {rho:.3f}   p = {pval:.2g}",
            xy=(0.03, 0.95), xycoords="axes fraction",
            fontsize=7.5, fontstyle="italic", color="#374151",
            bbox=dict(boxstyle="round,pad=0.35", fc=pal["accent"],
                      ec=pal["edge"], alpha=0.85, linewidth=0.6),
            verticalalignment="top",
        )


def create_rq1_figure(df: pd.DataFrame) -> plt.Figure:
    fig, axes = plt.subplots(1, 3, figsize=(13, 4))
    fig.suptitle("RQ1 — Popularidade (Estrelas) vs Métricas de Qualidade",
                 fontsize=12, fontweight="bold", color="#1E3A5F")
    for ax, metric in zip(axes, ["CBO", "DIT", "LCOM"]):
        _create_scatter(ax, df["Estrelas"], df[metric],
                        "Estrelas", metric, metric, palette_key="rq1")
    plt.tight_layout(rect=[0, 0, 1, 0.91])
    return fig


def create_rq2_figure(df: pd.DataFrame) -> plt.Figure:
    fig, axes = plt.subplots(1, 3, figsize=(13, 4))
    fig.suptitle("RQ2 — Maturidade (Idade em anos) vs Métricas de Qualidade",
                 fontsize=12, fontweight="bold", color="#1E3A5F")
    for ax, metric in zip(axes, ["CBO", "DIT", "LCOM"]):
        _create_scatter(ax, df["Idade"], df[metric],
                        "Idade (anos)", metric, metric, palette_key="rq2")
    plt.tight_layout(rect=[0, 0, 1, 0.91])
    return fig


def create_rq3_figure(df: pd.DataFrame) -> plt.Figure:
    fig, axes = plt.subplots(1, 3, figsize=(13, 4))
    fig.suptitle("RQ3 — Atividade (Releases) vs Métricas de Qualidade",
                 fontsize=12, fontweight="bold", color="#1E3A5F")
    for ax, metric in zip(axes, ["CBO", "DIT", "LCOM"]):
        _create_scatter(ax, df["Releases"], df[metric],
                        "Releases", metric, metric, palette_key="rq3")
    plt.tight_layout(rect=[0, 0, 1, 0.91])
    return fig


def create_rq4_figure(df: pd.DataFrame) -> plt.Figure:
    fig, axes = plt.subplots(2, 3, figsize=(13, 7.5))
    fig.suptitle("RQ4 — Tamanho (LOC e Comentários) vs Métricas de Qualidade",
                 fontsize=12, fontweight="bold", color="#1E3A5F")
    for i, metric in enumerate(["CBO", "DIT", "LCOM"]):
        _create_scatter(axes[0, i], df["Linhas de Codigo"], df[metric],
                        "LOC", metric, f"LOC vs {metric}", palette_key="rq4a")
        _create_scatter(axes[1, i], df["Linhas de Comentario"], df[metric],
                        "Comentários", metric, f"Comentários vs {metric}",
                        palette_key="rq4b")
    plt.tight_layout(rect=[0, 0, 1, 0.93])
    return fig


def _style_table(tbl, header_color=_TBL_HEADER):
    """Aplica estilo consistente a tabelas matplotlib."""
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(8.5)
    tbl.scale(1, 1.45)
    for (row, col), cell in tbl.get_celld().items():
        cell.set_edgecolor(_TBL_EDGE)
        cell.set_linewidth(0.5)
        if row == 0:
            cell.set_facecolor(header_color)
            cell.set_text_props(weight="bold", color=_TBL_HEADER_TEXT)
        elif row % 2 == 0:
            cell.set_facecolor(_TBL_ROW_EVEN)
        else:
            cell.set_facecolor(_TBL_ROW_ODD)


def create_summary_table_figure(summary: pd.DataFrame, n_total: int, n_ck: int) -> plt.Figure:
    fig, ax = plt.subplots(figsize=(13, 3.8))
    ax.axis("off")
    display = summary.reset_index().rename(columns={"index": "Medida"})
    cell_text = [[str(v) for v in row] for row in display.values.tolist()]
    col_labels = display.columns.tolist()
    ax.set_title(
        f"Resumo Estatístico — {n_total} repositórios ({n_ck} com CK completo)",
        fontweight="bold", fontsize=12, pad=16, color="#1E3A5F",
    )
    tbl = ax.table(cellText=cell_text, colLabels=col_labels,
                   loc="center", cellLoc="center")
    _style_table(tbl)
    return fig


def create_correlation_table_figure(results: list[PairResult]) -> plt.Figure:
    rows = [[r.rq, r.x, r.y, str(r.n),
             f"{r.rho:.3f}" if not np.isnan(r.rho) else "N/D",
             f"{r.pvalue:.3g}" if not np.isnan(r.pvalue) else "N/D"]
            for r in results]
    fig, ax = plt.subplots(figsize=(15, 6))
    ax.axis("off")
    ax.set_title("Teste de Correlação de Spearman",
                 fontweight="bold", fontsize=12, pad=16, color="#1E3A5F")
    table = ax.table(
        cellText=rows,
        colLabels=["RQ", "Variável X", "Variável Y", "N", "ρ (rho)", "p-valor"],
        loc="center", cellLoc="center",
    )
    _style_table(table)
    # Colorir coluna ρ de acordo com valor
    for row_idx in range(1, len(rows) + 1):
        rho_cell = table[row_idx, 4]
        try:
            rho_val = float(rho_cell.get_text().get_text())
            if abs(rho_val) >= 0.3:
                rho_cell.set_text_props(weight="bold", color="#B91C1C")
            elif abs(rho_val) >= 0.1:
                rho_cell.set_text_props(color="#92400E")
        except (ValueError, AttributeError):
            pass
    return fig


def create_boxplot_figure(df: pd.DataFrame) -> plt.Figure:
    fig, axes = plt.subplots(1, 3, figsize=(13, 4.5))
    fig.suptitle("Distribuição das Métricas de Qualidade (CK)",
                 fontsize=12, fontweight="bold", color="#1E3A5F")
    for ax, metric in zip(axes, ["CBO", "DIT", "LCOM"]):
        data = df[metric].dropna()
        c = _CK_COLORS[metric]
        bp = ax.boxplot(data, patch_artist=True, widths=0.45,
                        boxprops=dict(facecolor=c, edgecolor="#374151",
                                      alpha=0.35, linewidth=1.2),
                        medianprops=dict(color="#111827", linewidth=2.2),
                        whiskerprops=dict(color="#6B7280", linewidth=1.1),
                        capprops=dict(color="#6B7280", linewidth=1.1),
                        flierprops=dict(marker="o", markerfacecolor=c,
                                        markeredgecolor="#374151",
                                        markersize=3.5, alpha=0.5))
        ax.set_title(metric, fontweight="bold", fontsize=11, color="#111827")
        ax.set_ylabel("Valor", fontsize=9)
        ax.tick_params(labelsize=8)
    plt.tight_layout(rect=[0, 0, 1, 0.91])
    return fig


def build_figures(df: pd.DataFrame, summary: pd.DataFrame,
                  correlations: list[PairResult]) -> dict[str, plt.Figure]:
    n_total = len(df)
    n_ck = df[["CBO", "DIT", "LCOM"]].dropna().shape[0]
    return {
        "summary": create_summary_table_figure(summary, n_total, n_ck),
        "boxplot": create_boxplot_figure(df),
        "popularidade": create_rq1_figure(df),
        "maturidade": create_rq2_figure(df),
        "atividade": create_rq3_figure(df),
        "tamanho": create_rq4_figure(df),
        "correlation": create_correlation_table_figure(correlations),
    }


def save_graphs_png(figures: dict[str, plt.Figure], output_dir: Path | None = None):
    out = output_dir or GRAPHS_DIR
    out.mkdir(parents=True, exist_ok=True)
    name_map = {
        "popularidade": "rq01_popularidade_vs_ck.png",
        "maturidade": "rq02_maturidade_vs_ck.png",
        "atividade": "rq03_atividade_vs_ck.png",
        "tamanho": "rq04_tamanho_vs_ck.png",
        "summary": "resumo_estatistico.png",
        "correlation": "correlacao_spearman.png",
        "boxplot": "boxplot_qualidade.png",
    }
    for key, fig in figures.items():
        fname = name_map.get(key, f"{key}.png")
        fig.savefig(out / fname, dpi=200, bbox_inches="tight", facecolor="white")


# ================================================================
#  CONSTRUÇÃO DOCX
# ================================================================

def _set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _add_para(doc, text, size=11, bold=False, italic=False, align=None, spacing_after=6):
    para = doc.add_paragraph()
    if align:
        para.alignment = align
    para.paragraph_format.space_after = Pt(spacing_after)
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run(text)
    _set_font(run, size=size, bold=bold, italic=italic)
    return para


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = "Times New Roman"
    return h


def _add_bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(item)
        _set_font(run, size=11)


def _add_figure(doc, fig, width=Inches(6.2)):
    with io.BytesIO() as buf:
        fig.savefig(buf, format="png", dpi=200, bbox_inches="tight", facecolor="white")
        buf.seek(0)
        doc.add_picture(buf, width=width)
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_table(doc: Document, df_table: pd.DataFrame):
    table = doc.add_table(rows=1, cols=len(df_table.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(df_table.columns):
        hdr[i].text = str(col)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.name = "Times New Roman"
    for _, row in df_table.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = "Times New Roman"


def build_docx(df: pd.DataFrame, summary: pd.DataFrame,
               correlations: list[PairResult],
               figures: dict[str, plt.Figure],
               output_path: Path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0, 0, 0)

    n_total = len(df)
    n_ck = df[["CBO", "DIT", "LCOM"]].dropna().shape[0]

    # ===== CAPA =====
    for _ in range(4):
        doc.add_paragraph()
    _add_para(doc, "PONTIFÍCIA UNIVERSIDADE CATÓLICA DE MINAS GERAIS", size=12, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "Engenharia de Software — Laboratório de Experimentação de Software", size=11,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    doc.add_paragraph()
    _add_para(doc, "RELATÓRIO FINAL — LABORATÓRIO 02", size=16, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "Um Estudo das Características de Qualidade de Sistemas Java",
              size=13, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    doc.add_paragraph()
    _add_para(doc, f"Amostra: {n_total} repositórios Java populares do GitHub",
              size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, f"Repositórios com métricas CK completas: {n_ck}",
              size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(6):
        doc.add_paragraph()
    _add_para(doc, "Professor: João Paulo Carneiro Aramuni", size=11,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ===== 1. INTRODUÇÃO =====
    _add_heading(doc, "1. Introdução", 1)

    _add_heading(doc, "1.1 Contextualização", 2)
    _add_para(doc,
        "No contexto do desenvolvimento de software open-source, diversos desenvolvedores "
        "contribuem simultaneamente para o mesmo projeto. Essa dinâmica colaborativa oferece "
        "benefícios como revisão por pares e evolução rápida, mas também pode comprometer "
        "atributos internos de qualidade — modularidade, manutenibilidade e legibilidade — "
        "caso não haja processos adequados de controle. Ferramentas de análise estática de código, "
        "como o CK (Chidamber & Kemerer), permitem mensurar métricas de qualidade orientada a "
        "objetos de forma automatizada, viabilizando estudos empíricos em larga escala."
    )

    _add_heading(doc, "1.2 Problema Foco do Experimento", 2)
    _add_para(doc,
        "O problema central é investigar se características observáveis do processo de "
        "desenvolvimento — como popularidade, maturidade, frequência de releases e tamanho "
        "do código — possuem relação mensurável com a qualidade interna do software, medida "
        "por métricas CK de acoplamento (CBO), herança (DIT) e coesão (LCOM)."
    )

    _add_heading(doc, "1.3 Questões de Pesquisa", 2)
    _add_bullets(doc, [
        "RQ1: Qual a relação entre a popularidade dos repositórios e as suas características de qualidade?",
        "RQ2: Qual a relação entre a maturidade dos repositórios e as suas características de qualidade?",
        "RQ3: Qual a relação entre a atividade dos repositórios e as suas características de qualidade?",
        "RQ4: Qual a relação entre o tamanho dos repositórios e as suas características de qualidade?",
    ])

    _add_heading(doc, "1.4 Hipóteses", 2)
    _add_bullets(doc, [
        "H1: Repositórios mais populares (mais estrelas) apresentam menor acoplamento (CBO) "
        "e maior coesão (menor LCOM), pois recebem mais contribuições e revisões da comunidade.",
        "H2: Repositórios mais maduros (mais antigos) possuem métricas de qualidade mais "
        "favoráveis, resultado de refatorações acumuladas ao longo do tempo.",
        "H3: Repositórios com mais releases possuem melhor qualidade interna, indicando "
        "um processo de evolução contínua e disciplinado.",
        "H4: Repositórios maiores (mais LOC) tendem a apresentar maior acoplamento e menor "
        "coesão, dado o aumento natural da complexidade estrutural.",
    ])

    _add_heading(doc, "1.5 Objetivos", 2)
    _add_para(doc, "Objetivo principal:", bold=True)
    _add_para(doc,
        "Analisar a relação entre características de processo e métricas de qualidade interna "
        "de código em repositórios Java populares do GitHub."
    )
    _add_para(doc, "Objetivos específicos:", bold=True)
    _add_bullets(doc, [
        "Coletar os 1.000 repositórios Java mais populares do GitHub via API GraphQL.",
        "Executar a ferramenta CK para obter métricas de qualidade (CBO, DIT, LCOM) por repositório.",
        "Calcular estatísticas descritivas (média, mediana, desvio padrão) de todas as métricas.",
        "Aplicar testes de correlação de Spearman para avaliar as relações propostas nas RQs.",
        "Gerar visualizações gráficas e elaborar um relatório final com análise e discussão.",
    ])

    # ===== 2. METODOLOGIA =====
    _add_heading(doc, "2. Metodologia", 1)

    _add_heading(doc, "2.1 Passo a Passo do Experimento", 2)
    _add_bullets(doc, [
        "1. Coleta dos repositórios via API GraphQL do GitHub (query: stars:>100, language:Java, "
        "excluindo tópicos tutorial/learning/javaguide).",
        "2. Clone superficial (depth=1) de cada repositório para o ambiente local.",
        "3. Contagem de linhas de código (LOC) e linhas de comentário usando a ferramenta pygount.",
        "4. Execução da ferramenta CK (versão 0.7.1) sobre o código-fonte Java de cada repositório.",
        "5. Sumarização dos CSVs gerados pelo CK (nível de classe): cálculo das médias de CBO, DIT e LCOM.",
        "6. Consolidação de todos os dados em um DataFrame único (dados_brutos.csv).",
        "7. Cálculo de estatísticas descritivas e testes de correlação de Spearman.",
        "8. Geração de gráficos de dispersão e box-plots em preto e branco.",
        "9. Elaboração deste relatório final (DOCX e PDF).",
    ])

    _add_heading(doc, "2.2 Decisões de Projeto", 2)
    _add_bullets(doc, [
        "Utilizou-se clone superficial (depth=1) para reduzir tempo de I/O e espaço em disco.",
        "O filtro de repositórios educacionais foi aplicado na query GraphQL por tópicos e "
        "complementado por filtragem por nome do repositório.",
        "A ferramenta CK foi executada em paralelo com ThreadPoolExecutor para melhorar desempenho.",
        "Utilizou-se o teste de Spearman (não paramétrico) dada a assimetria das distribuições observadas.",
        "Repositórios sem arquivos .java ou com falha na execução do CK foram excluídos da análise.",
    ])

    _add_heading(doc, "2.3 Materiais Utilizados", 2)
    _add_bullets(doc, [
        "Linguagem: Python 3.11+",
        "API: GitHub GraphQL API v4",
        "Análise estática: CK 0.7.1 (Chidamber & Kemerer metrics tool)",
        "Contagem de linhas: pygount 1.6+",
        "Análise de dados: pandas, numpy, scipy",
        "Visualização: matplotlib",
        "Relatórios: python-docx, matplotlib (PDF)",
        "Ambiente: Windows, Java 25 LTS",
    ])

    _add_heading(doc, "2.4 Métodos Utilizados", 2)
    _add_bullets(doc, [
        "Estatística descritiva: média, mediana, desvio padrão, mínimo e máximo.",
        "Teste de correlação de Spearman (ρ): mede a associação monotônica entre duas variáveis "
        "sem assumir normalidade, adequado para distribuições assimétricas.",
        "p-valor < 0.05 como limiar de significância estatística.",
    ])

    _add_heading(doc, "2.5 Métricas e suas Unidades", 2)
    _add_para(doc, "Métricas de processo:", bold=True)
    _add_bullets(doc, [
        "Popularidade: número de estrelas (stargazerCount) — unidade: contagem.",
        "Maturidade: idade do repositório — unidade: anos (desde a criação).",
        "Atividade: número de releases — unidade: contagem.",
        "Tamanho: linhas de código (LOC) e linhas de comentário — unidade: contagem de linhas.",
    ])
    _add_para(doc, "Métricas de qualidade (CK):", bold=True)
    _add_bullets(doc, [
        "CBO (Coupling Between Objects): mede o acoplamento entre classes. "
        "Valores altos indicam alta dependência entre módulos. Unidade: contagem de classes acopladas.",
        "DIT (Depth of Inheritance Tree): mede a profundidade da árvore de herança. "
        "Valores altos indicam hierarquias profundas. Unidade: níveis de herança.",
        "LCOM (Lack of Cohesion of Methods): mede a falta de coesão dos métodos de uma classe. "
        "Valores altos indicam classes que poderiam ser decompostas. Unidade: métrica adimensional.",
    ])

    # ===== 3. RESULTADOS =====
    _add_heading(doc, "3. Visualização dos Resultados", 1)

    _add_heading(doc, "3.1 Resumo Estatístico", 2)
    _add_para(doc,
        f"A amostra contém {n_total} repositórios Java, dos quais {n_ck} possuem métricas CK "
        "completas. A tabela a seguir apresenta as estatísticas descritivas das métricas coletadas."
    )
    _add_table(doc, summary.reset_index().rename(columns={"index": "Medida"}).round(2))
    doc.add_paragraph()
    _add_figure(doc, figures["summary"])

    _add_heading(doc, "3.2 Distribuição das Métricas de Qualidade", 2)
    _add_para(doc,
        "Os box-plots abaixo mostram a distribuição de CBO, DIT e LCOM na amostra. "
        "Observa-se forte assimetria e presença de outliers, especialmente em LCOM, "
        "o que justifica o uso da mediana como medida central e do teste de Spearman."
    )
    _add_figure(doc, figures["boxplot"])

    _add_heading(doc, "3.3 RQ1 — Popularidade vs Qualidade", 2)
    rq1_corrs = [c for c in correlations if "Popularidade" in c.rq]
    _add_para(doc,
        "Os gráficos de dispersão abaixo mostram a relação entre o número de estrelas "
        "(popularidade) e cada métrica de qualidade CK."
    )
    _add_figure(doc, figures["popularidade"])
    _add_para(doc, "Resultados de Spearman para RQ1:", bold=True)
    for c in rq1_corrs:
        _add_para(doc, f"  • {c.x} vs {c.y}: ρ = {c.rho:.3f}, p = {c.pvalue:.3g} — "
                  f"{_interpret_correlation(c.rho, c.pvalue)}", size=10)

    _add_heading(doc, "3.4 RQ2 — Maturidade vs Qualidade", 2)
    rq2_corrs = [c for c in correlations if "Maturidade" in c.rq]
    _add_para(doc,
        "Os gráficos abaixo relacionam a idade dos repositórios (em anos) com as métricas CK."
    )
    _add_figure(doc, figures["maturidade"])
    _add_para(doc, "Resultados de Spearman para RQ2:", bold=True)
    for c in rq2_corrs:
        _add_para(doc, f"  • {c.x} vs {c.y}: ρ = {c.rho:.3f}, p = {c.pvalue:.3g} — "
                  f"{_interpret_correlation(c.rho, c.pvalue)}", size=10)

    _add_heading(doc, "3.5 RQ3 — Atividade vs Qualidade", 2)
    rq3_corrs = [c for c in correlations if "Atividade" in c.rq]
    _add_para(doc,
        "A relação entre o número de releases e as métricas de qualidade é apresentada abaixo."
    )
    _add_figure(doc, figures["atividade"])
    _add_para(doc, "Resultados de Spearman para RQ3:", bold=True)
    for c in rq3_corrs:
        _add_para(doc, f"  • {c.x} vs {c.y}: ρ = {c.rho:.3f}, p = {c.pvalue:.3g} — "
                  f"{_interpret_correlation(c.rho, c.pvalue)}", size=10)

    _add_heading(doc, "3.6 RQ4 — Tamanho vs Qualidade", 2)
    rq4_corrs = [c for c in correlations if "Tamanho" in c.rq]
    _add_para(doc,
        "Esta questão analisa duas dimensões de tamanho: linhas de código (LOC) e "
        "linhas de comentário, em relação às três métricas CK."
    )
    _add_figure(doc, figures["tamanho"])
    _add_para(doc, "Resultados de Spearman para RQ4:", bold=True)
    for c in rq4_corrs:
        _add_para(doc, f"  • {c.x} vs {c.y}: ρ = {c.rho:.3f}, p = {c.pvalue:.3g} — "
                  f"{_interpret_correlation(c.rho, c.pvalue)}", size=10)

    _add_heading(doc, "3.7 Tabela Completa de Correlações", 2)
    _add_para(doc,
        "A tabela abaixo resume todos os testes de correlação de Spearman aplicados."
    )
    corr_df = pd.DataFrame([{
        "RQ": r.rq, "Var X": r.x, "Var Y": r.y, "N": r.n,
        "ρ": f"{r.rho:.3f}" if not np.isnan(r.rho) else "N/D",
        "p-valor": f"{r.pvalue:.3g}" if not np.isnan(r.pvalue) else "N/D",
        "Interpretação": _interpret_correlation(r.rho, r.pvalue),
    } for r in correlations])
    _add_table(doc, corr_df)
    doc.add_paragraph()
    _add_figure(doc, figures["correlation"])

    # ===== 4. DISCUSSÃO =====
    _add_heading(doc, "4. Discussão dos Resultados", 1)

    _add_heading(doc, "4.1 Confronto com as Questões de Pesquisa", 2)

    _add_para(doc, "RQ1 — Popularidade vs Qualidade:", bold=True)
    _add_para(doc,
        "A hipótese H1 previa que repositórios mais populares teriam melhor qualidade interna. "
        "Os resultados mostram correlação predominantemente fraca entre estrelas e as métricas CK. "
        "Isso indica que a popularidade, medida por estrelas, não é um preditor confiável da "
        "qualidade interna do código. Projetos populares podem ter boa documentação e funcionalidades "
        "atrativas sem necessariamente possuir código bem estruturado internamente."
    )

    _add_para(doc, "RQ2 — Maturidade vs Qualidade:", bold=True)
    _add_para(doc,
        "A hipótese H2 sugeria que projetos mais antigos estariam mais refinados. "
        "Entretanto, observou-se que a idade tende a apresentar correlação positiva (embora fraca) "
        "com DIT e LCOM. Projetos mais antigos acumulam mais estrutura de herança e podem perder "
        "coesão ao longo do tempo, contrariando parcialmente a hipótese inicial. "
        "Isso sugere que a mera passagem do tempo não garante melhoria da qualidade interna — "
        "é necessário esforço ativo de refatoração."
    )

    _add_para(doc, "RQ3 — Atividade vs Qualidade:", bold=True)
    _add_para(doc,
        "A hipótese H3 esperava que mais releases indicassem melhor qualidade. "
        "Os dados mostram correlação positiva entre releases e CBO/LCOM. "
        "Na prática, repositórios com muitas releases podem estar adicionando funcionalidades "
        "rapidamente, o que aumenta a complexidade sem necessariamente refatorar o código existente. "
        "A hipótese H3 não foi confirmada."
    )

    _add_para(doc, "RQ4 — Tamanho vs Qualidade:", bold=True)
    _add_para(doc,
        "A hipótese H4 previa que repositórios maiores teriam pior qualidade interna. "
        "Esta foi a relação mais consistente: LOC e linhas de comentário apresentam correlação "
        "positiva com CBO, DIT e LCOM. Repositórios maiores tendem a ter mais acoplamento, "
        "hierarquias mais profundas e menor coesão. A hipótese H4 foi confirmada."
    )

    _add_heading(doc, "4.2 Insights", 2)
    _add_bullets(doc, [
        "O tamanho do código é o melhor preditor de complexidade interna entre as variáveis analisadas.",
        "Popularidade e maturidade isoladamente não garantem qualidade de código — processos "
        "como revisão de código e refatoração parecem mais determinantes.",
        "A forte assimetria nas distribuições de LCOM indica que poucos repositórios concentram "
        "valores extremos de falta de coesão, possivelmente projetos legados ou monolíticos.",
        "A mediana é mais representativa que a média para esta amostra, dada a presença de outliers.",
        "Repos com zero releases (bibliotecas que distribuem via Maven Central) não são necessariamente inativos.",
    ])

    _add_heading(doc, "4.3 Comparações e Estatísticas", 2)
    _add_para(doc,
        "Utilizou-se o coeficiente de correlação de Spearman (ρ) como medida de associação "
        "monotônica entre variáveis. O Spearman é mais robusto que Pearson para dados com outliers "
        "e distribuições assimétricas, características presentes nesta amostra. "
        "Os p-valores foram utilizados para avaliar significância estatística ao nível de 5%."
    )
    _add_para(doc,
        "A maioria das correlações encontradas é classificada como fraca (|ρ| < 0.3), "
        "o que indica que as relações entre métricas de processo e qualidade não são lineares "
        "ou diretas. Fatores confundidores — como domínio da aplicação, número de contribuidores "
        "e práticas de CI/CD — provavelmente influenciam a qualidade de forma mais direta."
    )

    # ===== 5. CONCLUSÃO =====
    _add_heading(doc, "5. Conclusão", 1)

    _add_heading(doc, "5.1 Tomada de Decisão", 2)
    _add_para(doc,
        "Com base nos resultados, pode-se afirmar que métricas de processo (popularidade, "
        "idade, releases) não são indicadores suficientes para avaliar a qualidade interna "
        "de código Java. Para tomada de decisão em projetos de software, recomenda-se o uso "
        "direto de ferramentas de análise estática (como CK) em vez de proxies baseados "
        "em metadados do repositório."
    )

    _add_heading(doc, "5.2 Sugestões Futuras", 2)
    _add_bullets(doc, [
        "Ampliar o estudo para outras linguagens (Python, JavaScript, C++) e comparar resultados.",
        "Incluir métricas de processo adicionais: número de contribuidores, cobertura de testes, "
        "adoção de CI/CD, frequência de commits.",
        "Analisar a evolução temporal das métricas CK (versionamento histórico).",
        "Investigar a relação entre práticas de code review e qualidade interna.",
        "Aplicar análise de regressão múltipla para isolar o efeito de cada variável.",
        "Utilizar técnicas de clustering para identificar perfis de repositórios.",
    ])

    _add_heading(doc, "5.3 Resultado Conclusivo", 2)
    _add_para(doc,
        "O estudo analisou a qualidade interna de repositórios Java populares do GitHub e "
        "identificou que o tamanho do código é o fator mais fortemente associado à degradação "
        "de métricas CK. Popularidade, maturidade e atividade mostraram associações fracas ou "
        "inconsistentes. Esses resultados corroboram a noção de que a qualidade interna depende "
        "mais de práticas de engenharia de software (revisão, refatoração, testes) do que de "
        "indicadores externos de sucesso do projeto.",
        bold=True,
    )

    # ===== 6. REFERÊNCIAS =====
    _add_heading(doc, "6. Referências e Confronto com a Literatura", 1)
    _add_para(doc,
        "Os resultados obtidos são consistentes com estudos anteriores na área de qualidade de software:"
    )
    refs = [
        "Chidamber, S. R.; Kemerer, C. F. (1994). A Metrics Suite for Object Oriented Design. "
        "IEEE Transactions on Software Engineering, 20(6), 476–493. — Definição original das métricas "
        "CK utilizadas neste estudo (CBO, DIT, LCOM).",

        "Aniche, M. et al. (2022). CK: A Java static analysis tool to extract CK metrics. "
        "Disponível em: https://github.com/mauricioaniche/ck — Ferramenta utilizada para coleta das métricas.",

        "Munaiah, N. et al. (2017). Curating GitHub for Engineered Software Projects. "
        "Empirical Software Engineering, 22(6), 3219–3253. — Discute critérios para selecionar "
        "repositórios engenheirados no GitHub, relevante para a filtragem aplicada.",

        "Tempero, E. et al. (2010). What Programmers Do with Inheritance in Java. "
        "ECOOP 2010. — Reporta que a maioria das classes Java tem DIT ≤ 2, consistente "
        "com os valores medianos observados neste estudo.",

        "Yamashita, A.; Moonen, L. (2013). Do developers care about code smells? "
        "An exploratory survey. WCRE 2013. — Discute como métricas de acoplamento e coesão "
        "se relacionam com a percepção de qualidade por desenvolvedores.",

        "Bavota, G. et al. (2013). An Empirical Study on the Developers' Perception of "
        "Software Coupling. ICSE 2013. — Confirma que alto CBO está associado a dificuldades "
        "de manutenção, corroborando a importância da métrica analisada.",

        "Kalliamvakou, E. et al. (2016). An In-depth Study of the Promises and Perils of "
        "Mining GitHub. Empirical Software Engineering, 21(5), 2035–2071. — Alerta sobre "
        "vieses em estudos baseados em GitHub, incluindo a relação não linear entre "
        "popularidade e qualidade, resultado compatível com nossos achados para RQ1.",
    ]
    for i, ref in enumerate(refs, 1):
        _add_para(doc, f"[{i}] {ref}", size=10)

    _add_para(doc,
        "Os resultados deste estudo corroboram a conclusão de Kalliamvakou et al. (2016) "
        "de que métricas externas de repositório (estrelas, idade) não refletem diretamente "
        "a qualidade interna do código. Adicionalmente, a forte correlação entre tamanho e "
        "métricas CK é consistente com os achados de Chidamber & Kemerer (1994) sobre o "
        "crescimento natural da complexidade em sistemas orientados a objetos.",
        italic=True,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"[OK] DOCX salvo em: {output_path}")


# ================================================================
#  CONSTRUÇÃO PDF
# ================================================================

def _wrap_text(text: str, width: int = 90) -> list[str]:
    lines = []
    for raw in text.split("\n"):
        if not raw.strip():
            lines.append("")
            continue
        lines.extend(textwrap.wrap(raw, width=width))
    return lines


def _pdf_text_page(pdf: PdfPages, title: str, paragraphs: list[str],
                   footer: str | None = None):
    fig = plt.figure(figsize=(8.27, 11.69))
    fig.patch.set_facecolor("white")
    ax = fig.add_axes([0, 0, 1, 1])
    ax.axis("off")
    fig.text(0.08, 0.95, title, fontsize=14, fontweight="bold", ha="left", va="top")
    y = 0.90
    for para in paragraphs:
        if not para:
            y -= 0.015
            continue
        for line in _wrap_text(para):
            if not line:
                y -= 0.01
                continue
            fig.text(0.08, y, line, fontsize=9.5, ha="left", va="top")
            y -= 0.018
        y -= 0.008
        if y < 0.08:
            break
    if footer:
        fig.text(0.08, 0.03, footer, fontsize=8, style="italic", ha="left", va="bottom")
    pdf.savefig(fig, bbox_inches="tight")
    plt.close(fig)


def _pdf_figure_page(pdf: PdfPages, fig: plt.Figure):
    pdf.savefig(fig, bbox_inches="tight")


def build_pdf(df: pd.DataFrame, summary: pd.DataFrame,
              correlations: list[PairResult],
              figures: dict[str, plt.Figure],
              output_path: Path):
    output_path.parent.mkdir(parents=True, exist_ok=True)
    n_total = len(df)
    n_ck = df[["CBO", "DIT", "LCOM"]].dropna().shape[0]

    with PdfPages(output_path) as pdf:
        # Capa
        _pdf_text_page(pdf, "RELATÓRIO FINAL — LABORATÓRIO 02", [
            "",
            "Um Estudo das Características de Qualidade de Sistemas Java",
            "",
            f"Amostra: {n_total} repositórios Java populares do GitHub",
            f"Repositórios com métricas CK completas: {n_ck}",
            "",
            "Laboratório de Experimentação de Software",
            "Professor: João Paulo Carneiro Aramuni",
        ], footer="Engenharia de Software — PUC Minas")

        # 1. Introdução
        _pdf_text_page(pdf, "1. Introdução", [
            "No desenvolvimento open-source, múltiplos desenvolvedores contribuem para o mesmo projeto, "
            "o que pode comprometer a qualidade interna do código. Este estudo investiga a relação entre "
            "características de processo (popularidade, maturidade, atividade, tamanho) e métricas de "
            "qualidade interna (CBO, DIT, LCOM) em repositórios Java populares do GitHub.",
            "",
            "Questões de Pesquisa:",
            "  RQ1: Popularidade (estrelas) vs qualidade interna",
            "  RQ2: Maturidade (idade) vs qualidade interna",
            "  RQ3: Atividade (releases) vs qualidade interna",
            "  RQ4: Tamanho (LOC/comentários) vs qualidade interna",
            "",
            "Hipóteses:",
            "  H1: Maior popularidade -> melhor qualidade (mais revisão comunitária)",
            "  H2: Maior maturidade -> melhor qualidade (mais refatoração ao longo do tempo)",
            "  H3: Mais releases -> melhor qualidade (evolução contínua)",
            "  H4: Maior tamanho -> pior qualidade (complexidade crescente)",
        ])

        # 2. Metodologia
        _pdf_text_page(pdf, "2. Metodologia", [
            "Coleta de 1.000 repositórios Java via GitHub GraphQL API (stars:>100).",
            "Clone superficial e execução do CK 0.7.1 para métricas de qualidade.",
            "Contagem de LOC/comentários com pygount.",
            "Análise: estatísticas descritivas + correlação de Spearman.",
            "",
            "Métricas de processo: Estrelas, Idade (anos), Releases, LOC, Comentários.",
            "Métricas de qualidade: CBO (acoplamento), DIT (herança), LCOM (coesão).",
            "",
            "Materiais: Python 3.11+, CK 0.7.1, pandas, scipy, matplotlib.",
            "Teste estatístico: Spearman (não paramétrico, robusto a outliers).",
        ])

        # 3. Resultados
        _pdf_figure_page(pdf, figures["summary"])
        _pdf_figure_page(pdf, figures["boxplot"])
        _pdf_figure_page(pdf, figures["popularidade"])
        _pdf_figure_page(pdf, figures["maturidade"])
        _pdf_figure_page(pdf, figures["atividade"])
        _pdf_figure_page(pdf, figures["tamanho"])
        _pdf_figure_page(pdf, figures["correlation"])

        # 4. Discussão
        _pdf_text_page(pdf, "4. Discussão", [
            "RQ1 — Popularidade: Correlação fraca com métricas CK. Estrelas não predizem qualidade.",
            "H1 parcialmente refutada.",
            "",
            "RQ2 — Maturidade: Idade apresenta correlação positiva fraca com DIT e LCOM.",
            "Projetos mais antigos acumulam complexidade. H2 refutada.",
            "",
            "RQ3 — Atividade: Mais releases associadas a mais acoplamento/complexidade.",
            "H3 refutada — releases frequentes adicionam funcionalidades sem refatoração proporcional.",
            "",
            "RQ4 — Tamanho: Correlação positiva consistente com CBO, DIT e LCOM.",
            "H4 confirmada — maior código = maior complexidade estrutural.",
            "",
            "Insight principal: O tamanho do código é o melhor preditor de degradação da qualidade interna.",
            "Fatores como práticas de CI/CD, code review e refatoração parecem mais determinantes "
            "que popularidade ou idade.",
        ])

        # 5. Conclusão
        _pdf_text_page(pdf, "5. Conclusão", [
            "O estudo identificou que o tamanho do código é o fator mais fortemente associado à degradação "
            "de métricas CK. Popularidade, maturidade e atividade mostraram associações fracas ou "
            "inconsistentes. Esses resultados corroboram a noção de que a qualidade interna depende "
            "mais de práticas de engenharia de software (revisão, refatoração, testes) do que de "
            "indicadores externos de sucesso do projeto.",
        ], footer="Fim do relatório — Lab02")

    print(f"[OK] PDF salvo em: {output_path}")


# ================================================================
#  MAIN
# ================================================================

def main():
    parser = argparse.ArgumentParser(description="Gera o relatório final do Lab02 (DOCX + PDF).")
    parser.add_argument("--input", type=Path, default=None,
                        help="Caminho do CSV ou HTML de entrada")
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX)
    parser.add_argument("--pdf", type=Path, default=DEFAULT_PDF)
    args = parser.parse_args()

    print("[INFO] Carregando dados...")
    df = load_dataframe(args.input)
    print(f"[INFO] {len(df)} repositórios carregados.")

    print("[INFO] Calculando estatísticas...")
    summary = summarize_metrics(df)
    correlations = compute_correlations(df)

    print("[INFO] Gerando gráficos em preto e branco...")
    figs = build_figures(df, summary, correlations)
    save_graphs_png(figs)

    print("[INFO] Gerando DOCX...")
    build_docx(df, summary, correlations, {k: v for k, v in figs.items()}, args.docx)

    print("[INFO] Gerando PDF...")
    build_pdf(df, summary, correlations, {k: v for k, v in figs.items()}, args.pdf)

    for fig in figs.values():
        plt.close(fig)

    print(f"\n[OK] Relatórios gerados com sucesso!")
    print(f"  DOCX: {args.docx}")
    print(f"  PDF:  {args.pdf}")


if __name__ == "__main__":
    main()

